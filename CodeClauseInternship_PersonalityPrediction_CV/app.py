
# import os
# import PyPDF2
# import docx
# from flask import Flask, render_template, request, redirect, url_for, flash
# from werkzeug.utils import secure_filename
# import pandas as pd
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.ensemble import RandomForestRegressor
# from sklearn.model_selection import train_test_split

# app = Flask(__name__)
# app.secret_key = 'secretkey'

# UPLOAD_FOLDER = 'static/uploads'
# if not os.path.exists(UPLOAD_FOLDER):
#     os.makedirs(UPLOAD_FOLDER)
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# file_path = 'PersonalityPrediction\\archive\\IPIP-FFM-data-8Nov2018\\mypersonality_final.csv'
# df = pd.read_csv(file_path, encoding='ISO-8859-1')

# text_column = 'STATUS'  # Text data (resume content or posts)
# target_columns = ['sEXT', 'sNEU', 'sAGR', 'sCON', 'sOPN']  # Big Five traits

# df = df.dropna(subset=[text_column] + target_columns).head(1000)

# tfidf_vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
# X = tfidf_vectorizer.fit_transform(df[text_column]).toarray()
# y = df[target_columns].values

# X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
# rf_model = RandomForestRegressor(n_estimators=50, random_state=42)
# rf_model.fit(X_train, y_train)

# def extract_text_from_pdf(file_path):
#     with open(file_path, 'rb') as f:
#         reader = PyPDF2.PdfReader(f)
#         text = ""
#         for page in reader.pages:
#             text += page.extract_text() or ""
#     return text

# def extract_text_from_docx(file_path):
#     doc = docx.Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/upload', methods=['POST'])
# def upload_cv():
#     if 'file' not in request.files:
#         flash('No file part')
#         return redirect(request.url)

#     file = request.files['file']
#     if file.filename == '':
#         flash('No selected file')
#         return redirect(request.url)

#     if file:
#         filename = secure_filename(file.filename)
#         filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         file.save(filepath)

#         if filename.endswith('.pdf'):
#             cv_text = extract_text_from_pdf(filepath)
#         elif filename.endswith('.docx'):
#             cv_text = extract_text_from_docx(filepath)
#         else:
#             flash('Unsupported file format. Please upload a PDF or DOCX file.')
#             return redirect(request.url)

#         if not cv_text.strip():
#             flash('The uploaded file is empty or could not be processed.')
#             return redirect(request.url)

#         cv_tfidf = tfidf_vectorizer.transform([cv_text]).toarray()
#         predicted_traits = rf_model.predict(cv_tfidf)[0]

#         traits_summary = {
#             "Extroversion": round(predicted_traits[0], 2),
#             "Neuroticism": round(predicted_traits[1], 2),
#             "Agreeableness": round(predicted_traits[2], 2),
#             "Conscientiousness": round(predicted_traits[3], 2),
#             "Openness": round(predicted_traits[4], 2)
#         }

#         return render_template('result.html', traits=traits_summary, filename=filename)

# if __name__ == "__main__":
#     app.run(debug=True, use_reloader=False)
import os
import PyPDF2
import docx
import matplotlib.pyplot as plt
from flask import Flask, render_template, request, redirect, url_for, flash
from werkzeug.utils import secure_filename
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split
from io import BytesIO
import base64
import re

from flask import Flask, render_template, request, redirect, url_for, flash
from werkzeug.utils import secure_filename
import os
import PyPDF2
import docx
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split
import re

app = Flask(__name__)
app.secret_key = 'secretkey'
UPLOAD_FOLDER = 'static/uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

file_path = 'PersonalityPrediction\\archive\\IPIP-FFM-data-8Nov2018\\mypersonality_final.csv'
df = pd.read_csv(file_path, encoding='ISO-8859-1')
text_column = 'STATUS'  
target_columns = ['sEXT', 'sNEU', 'sAGR', 'sCON', 'sOPN'] 
df = df.dropna(subset=[text_column] + target_columns).head(1000)

tfidf_vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
X = tfidf_vectorizer.fit_transform(df[text_column]).toarray()
y = df[target_columns].values

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
rf_model = RandomForestRegressor(n_estimators=50, random_state=42)
rf_model.fit(X_train, y_train)

general_keywords = ["experience", "skills", "education", "projects", "achievements", "responsibilities", "certification"]

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def calculate_general_ats_score(cv_text):
    keyword_count = sum(cv_text.lower().count(word) for word in general_keywords)
    keyword_score = (keyword_count / len(general_keywords)) * 20  

    sections = ["experience", "education", "skills"]
    section_score = sum(1 for section in sections if re.search(rf"\b{section}\b", cv_text, re.IGNORECASE)) * 20  

    formatting_score = 20 if len(cv_text) > 500 else 0  

    ats_score = keyword_score + section_score + formatting_score
    return round(ats_score, 2)

def generate_plot(traits, ats_score):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))

    labels = ["Extroversion", "Neuroticism", "Agreeableness", "Conscientiousness", "Openness"]
    values = [traits[label] for label in labels]
    ax1.bar(labels, values, color=['blue', 'red', 'green', 'orange', 'purple'])
    ax1.set_ylim(0, 100)
    ax1.set_title("Personality Traits")
    ax1.set_ylabel("Percentage")

    ax2.bar(["ATS Score"], [ats_score], color='skyblue')
    ax2.set_ylim(0, 100)
    ax2.set_title("General ATS Compatibility Score")
    ax2.set_ylabel("Score (%)")

    buf = BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png")
    buf.seek(0)
    image_base64 = base64.b64encode(buf.read()).decode("utf-8")
    buf.close()
    plt.close(fig)
    return image_base64

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_cv():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(request.url)

    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect(request.url)

    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        if filename.endswith('.pdf'):
            cv_text = extract_text_from_pdf(filepath)
        elif filename.endswith('.docx'):
            cv_text = extract_text_from_docx(filepath)
        else:
            flash('Unsupported file format. Please upload a PDF or DOCX file.')
            return redirect(request.url)

        if not cv_text.strip():
            flash('The uploaded file is empty or could not be processed.')
            return redirect(request.url)
        cv_tfidf = tfidf_vectorizer.transform([cv_text]).toarray()
        predicted_traits = rf_model.predict(cv_tfidf)[0]
        traits_summary = {
            "Extroversion": round(predicted_traits[0] * 100, 2),
            "Neuroticism": round(predicted_traits[1] * 100, 2),
            "Agreeableness": round(predicted_traits[2] * 100, 2),
            "Conscientiousness": round(predicted_traits[3] * 100, 2),
            "Openness": round(predicted_traits[4] * 100, 2)
        }


        ats_score = calculate_general_ats_score(cv_text)


        return render_template('result.html', traits=traits_summary, ats_score=ats_score, filename=filename)

if __name__ == "__main__":
    app.run(debug=True, use_reloader=False)
